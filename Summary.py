import os

import docx

import openai


def generate_summary(input_doc_path, output_dir):
    # 设置 Azure OpenAI API 参数
    openai.api_type = "azure"
    openai.api_base = "https://cshitinternopenai.openai.azure.com/"
    openai.api_version = "2023-03-15-preview"
    openai.api_key = "0be4adcd512d4b09b7e44d50325f4bf9"

    # 读取输入文档
    # input_doc_path = "C:\\users\\user\\Desktop\\112\\20230707-1\\音檔轉文字\\OutputWord\\output.docx"  # 替换为实际的输入文档路径
    doc = docx.Document(input_doc_path)

    # 提取文档内容
    content = ""
    for paragraph in doc.paragraphs:
        content += paragraph.text + "\n"

    # 使用 Azure OpenAI API 生成文章摘要
    response = openai.ChatCompletion.create(
        engine="CSHITIntern",  # engine = "deployment_name".
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": content},
        ]
    )

    summary = response.choices[0].message.content

    # 创建和保存输出文档
    # output_dir = "C:\\users\\user\\Desktop\\112\\20230707-1\\音檔轉文字"  # 替换为实际的输出目录路径
    os.makedirs(output_dir, exist_ok=True)

    final_dir = os.path.join(output_dir, "Final")
    os.makedirs(final_dir, exist_ok=True)

    output_doc_path = os.path.join(final_dir, "Final.docx")
    output_doc = docx.Document()
    output_doc.add_paragraph(summary)
    output_doc.save(output_doc_path)

    # print("摘要生成并保存成功。")
