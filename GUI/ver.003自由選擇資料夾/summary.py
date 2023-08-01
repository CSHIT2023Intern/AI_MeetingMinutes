import os

import docx

import openai


def generate_summary(input_doc_path, output_dir):
    # Azure OpenAI API
    openai.api_type = "azure"
    openai.api_base = "https://cshitinternopenai.openai.azure.com/"
    openai.api_version = "2023-03-15-preview"
    openai.api_key = "0be4adcd512d4b09b7e44d50325f4bf9"

    # 輸入檔案
    doc = docx.Document(input_doc_path)

    # 存文件
    content = ""
    for paragraph in doc.paragraphs:
        content += paragraph.text + "\n"

    # 生成摘要
    response = openai.ChatCompletion.create(
        engine="CSHITIntern",
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": content},
            {"role": "assistant", "content": "幫我進行摘要"}
        ]
    )

    summary = response.choices[0].message.content

    # 存檔
    final_dir = os.path.join(output_dir, "會議摘要")
    os.makedirs(final_dir, exist_ok=True)

    output_doc_path = os.path.join(final_dir, "會議摘要.docx")
    output_doc = docx.Document()
    output_doc.add_paragraph(summary)
    output_doc.save(output_doc_path)
