# 合體方法
import os

import docx

import openai

import os
import azure.cognitiveservices.speech as speechsdk
from docx import Document

SPEECH_KEY = ""
SPEECH_REGION = ""


def trsns(filename, output_filepath):
    speech_config = speechsdk.SpeechConfig(
        subscription=SPEECH_KEY, region=SPEECH_REGION)
    speech_config.speech_recognition_language = "zh-TW"
    audio_config = speechsdk.audio.AudioConfig(filename=filename)
    speech_recognizer = speechsdk.SpeechRecognizer(
        speech_config=speech_config, audio_config=audio_config)

    speech_recognition_result = speech_recognizer.recognize_once_async().get()
    recognized_text = speech_recognition_result.text

    # 存到WORD
    doc = Document()
    doc.add_paragraph(recognized_text)
    doc.save(output_filepath)


def generate_summary(input_doc_path, output_dir):
    # Azure OpenAI API
    openai.api_type = "azure"
    openai.api_base = "https://cshitinternopenai.openai.azure.com/"
    openai.api_version = "2023-03-15-preview"
    openai.api_key = "0be4adcd512d4b09b7e44d50325f4bf9"

    # 輸入檔案
    doc = docx.Document(input_doc_path)

    # 存文件
    content = ""
    for paragraph in doc.paragraphs:
        content += paragraph.text + "\n"

    # 生成摘要
    response = openai.ChatCompletion.create(
        engine="CSHITIntern",
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": content},
            {"role": "assistant", "content": "幫我進行摘要"}
        ]
    )

    summary = response.choices[0].message.content

    # 存檔
    final_dir = os.path.join(output_dir, "會議摘要")
    os.makedirs(final_dir, exist_ok=True)

    output_doc_path = os.path.join(final_dir, "會議摘要.docx")
    output_doc = docx.Document()
    output_doc.add_paragraph(summary)
    output_doc.save(output_doc_path)


input_filename = "/content/test.wav"
output_filename = "逐字稿.docx"
trsns(input_filename, output_filename)
output_dir = "/content"
input_doc_path = output_filename
generate_summary(input_doc_path, output_dir)
